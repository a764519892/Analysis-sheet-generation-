from docx import Document
from docx.shared import Inches, Pt,Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time,datetime
from docx.enum.table import WD_ALIGN_VERTICAL,WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import time
from os import walk
def yun_xing(url):
    # name = input('姓名：')
    # departments = input('院系')
    # specialty = input('专业')
    # class_name = input('班级：')
    # gender = input('性别：')
    # school_num = input('学号：')
    # school_time_year = input('在校起始时间年')
    # school_time_month = input('在校起始时间月')
    # school_time_year1 = input('在校截止时间年')
    # school_time_month1 = input('在校截止时间月')
    # home_address = input('家庭通讯地址：')
    # personal_tel = input('个人联系方式：')
    # home_num = input('家庭联系方式：')
    # reason = input('本人申请理由：')

    nt = datetime.datetime.now()
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    p = document.add_paragraph()
    p1 = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # 设置段落从右开始缩进，使用Pt来衡量
    p1.paragraph_format.right_indent = Pt(20)
    r = p.add_run('香港瑞升（越南）纺织品有限公司\r来 样 分 析 登 记 表')
    r.font.size = Pt(16)
    r.bold = True
    directory = url
    t = 0
    for root, dirs, filenames in walk(directory):
        for file in filenames:
            if file.endswith(".doc") or file.endswith(".docx"):
                t = t + 1
    bianma = t + 1
    ri_qi = str(time.strftime("%Y%m%d", time.localtime()))
    deng_ji_hao = ri_qi[2:] + '-%s' % bianma
    r1 = p1.add_run('登记号：{}'.format(deng_ji_hao))
    r1.font.size = Pt(10)
    r1.bold = True

    table = document.add_table(rows=37, cols=16, style='Table Grid')
    table.style.font.name = 'Arial'
    table.style.font.size = Pt(10)
    # run = table.cell(0,0).paragraphs[0].add_run('牛逼')
    # run.font.name = 'Arial'
    # run.font.size = 140000
    table.autofit = False
    table.columns[0].width = Inches(0.49)  # 0.49
    # table.columns[2].width = Inches(0.85) #0.49
    # table.columns[4].width = Inches(0.85) #0.49
    # table.columns[6].width = Inches(0.85) #0.49
    # table.columns[8].width = Inches(0.85) #0.49
    # 第一排
    table.cell(0, 0).merge(table.cell(1, 1))
    table.cell(0, 2).merge(table.cell(1, 3))
    table.cell(0, 4).merge(table.cell(1, 5))
    table.cell(0, 6).merge(table.cell(1, 7))
    table.cell(0, 8).merge(table.cell(1, 9))
    table.cell(0, 10).merge(table.cell(1, 11))
    table.cell(0, 12).merge(table.cell(1, 13))
    table.cell(0, 14).merge(table.cell(1, 15))
    # 第二排
    table.cell(2, 0).merge(table.cell(3, 1))
    table.cell(2, 2).merge(table.cell(3, 3))
    table.cell(2, 4).merge(table.cell(3, 5))
    table.cell(2, 6).merge(table.cell(3, 7))
    table.cell(2, 8).merge(table.cell(3, 9))
    table.cell(2, 10).merge(table.cell(3, 11))
    table.cell(2, 12).merge(table.cell(3, 13))
    table.cell(2, 14).merge(table.cell(3, 15))
    # 第三排
    table.cell(4, 0).merge(table.cell(5, 1))
    table.cell(4, 2).merge(table.cell(5, 3))
    table.cell(4, 4).merge(table.cell(5, 5))
    table.cell(4, 6).merge(table.cell(5, 7))
    table.cell(4, 8).merge(table.cell(5, 9))
    table.cell(4, 10).merge(table.cell(5, 11))
    table.cell(4, 12).merge(table.cell(5, 13))
    table.cell(4, 14).merge(table.cell(5, 15))
    # 第四排
    table.cell(6, 0).merge(table.cell(7, 1))
    table.cell(6, 2).merge(table.cell(7, 15))
    # 第五排
    table.cell(8, 0).merge(table.cell(9, 1))
    table.cell(8, 2).merge(table.cell(9, 15))
    # 第6
    table.cell(10, 0).merge(table.cell(11, 1))
    table.cell(10, 2).merge(table.cell(11, 15))
    # 第7
    table.cell(12, 0).merge(table.cell(15, 1))
    table.cell(12, 2).merge(table.cell(15, 15))
    # 第8
    table.cell(14, 0).merge(table.cell(36, 15))
    hdr_cells0 = table.rows[0].cells
    hdr_cells2 = table.rows[2].cells
    hdr_cells4 = table.rows[4].cells
    hdr_cells6 = table.rows[6].cells
    hdr_cells8 = table.rows[8].cells
    hdr_cells10 = table.rows[10].cells
    # hdr_cells21 = table.rows[21].cells
    # hdr_cells27 = table.rows[27].cells
    # hdr_cells32 = table.rows[32].cells

    # hdr_cells0[0].add_paragraph('客户名称').alignment=WD_ALIGN_PARAGRAPH.CENTER #居中
    # hdr_cells0[0].add_paragraph('客户名称').alignment=WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    # table.rows[0].height=Cm(0.30)
    # 添加标签
    table.cell(0, 0).paragraphs[0].add_run('客户名称')
    table.cell(0, 4).paragraphs[0].add_run('订单号')
    table.cell(0, 8).paragraphs[0].add_run('面料名称')
    table.cell(0, 12).paragraphs[0].add_run('颜色')

    table.cell(2, 0).paragraphs[0].add_run('要求克重')
    table.cell(2, 4).paragraphs[0].add_run('实际克重')
    table.cell(2, 8).paragraphs[0].add_run('要求成分')
    table.cell(2, 12).paragraphs[0].add_run('实测成分')

    table.cell(4, 0).paragraphs[0].add_run('织造设备')
    table.cell(4, 4).paragraphs[0].add_run('机号（GN）')
    table.cell(4, 8).paragraphs[0].add_run('特殊加工')
    table.cell(4, 12).paragraphs[0].add_run('门幅')

    table.cell(6, 0).paragraphs[0].add_run('纱支分析')

    table.cell(8, 0).paragraphs[0].add_run('纱长50C')

    table.cell(10, 0).paragraphs[0].add_run('纱比%')
    # 读取内容
    with open(r'D:\Study\pythonProject\remi_design\jisuan_shabu_canshu\ceshi\1.txt', 'r', encoding='utf-8') as f:
        neirong = f.read().split('\n')
        neirong.pop()
    # 添加内容

    table.cell(0, 2).paragraphs[0].add_run('')
    table.cell(0, 6).paragraphs[0].add_run('')
    table.cell(0, 10).paragraphs[0].add_run('')
    table.cell(0, 14).paragraphs[0].add_run('')

    table.cell(2, 2).paragraphs[0].add_run('')
    table.cell(2, 6).paragraphs[0].add_run(neirong[0])
    table.cell(2, 10).paragraphs[0].add_run('')
    table.cell(2, 14).paragraphs[0].add_run('')

    table.cell(4, 2).paragraphs[0].add_run('')
    table.cell(4, 6).paragraphs[0].add_run(neirong[1])
    table.cell(4, 10).paragraphs[0].add_run('')
    table.cell(4, 14).paragraphs[0].add_run(neirong[2])

    table.cell(6, 2).paragraphs[0].add_run(neirong[3])

    table.cell(8, 2).paragraphs[0].add_run(neirong[4])

    table.cell(10, 2).paragraphs[0].add_run('')
    # table.cell(12,0).paragraphs[0].add_run('客户要求')

    # hdr_cells0[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells0[0].add_paragraph('牛逼').vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    # table.cell(0,1).add_paragraph('客户名称').vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # 两个想加 真正文字居中
    #table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    #table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ##单元格垂直居中
    #太慢，每个都居中
    for row in range(0,5,2):
        for col in range(0,15,2):
            table.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
            table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ##单元格垂直居中
    table.cell(6, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    table.cell(6, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ##单元格垂直居中
    table.cell(8, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    table.cell(8, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ##单元格垂直居中
    table.cell(6, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    table.cell(6, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ##单元格垂直居中
    table.cell(8, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    table.cell(8, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ##单元格垂直居中
    table.cell(10, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    table.cell(10, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ##单元格垂直居中
    table.cell(10, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    table.cell(10, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ##单元格垂直居中
    table.cell(12, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中
    table.cell(12, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER  ##单元格垂直居中
    #给需要居中


    # hdr_cells0[4].add_paragraph('订单号').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells0[8].add_paragraph('面料名称').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells0[12].add_paragraph('颜色').alignment=WD_ALIGN_PARAGRAPH.CENTER
    #
    # hdr_cells3[0].add_paragraph('姓名\n').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells3[3].add_paragraph(name).alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells3[6].add_paragraph('性别').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells3[7].add_paragraph(gender).alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells3[10].add_paragraph('学号').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells3[11].add_paragraph(school_num).alignment=WD_ALIGN_PARAGRAPH.CENTER
    #
    # hdr_cells6[0].add_paragraph('在校时间\n').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells6[3].add_paragraph(school_time_year + '年' + school_time_month + '月——' + school_time_year1 + '年'+ school_time_month1 + '月').alignment=WD_ALIGN_PARAGRAPH.CENTER
    #
    # hdr_cells9[0].add_paragraph('家庭通讯地址').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells9[3].add_paragraph(home_address).alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells9[7].add_paragraph('家庭联系方式').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells12[7].add_paragraph('个人联系方式').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells9[10].add_paragraph(home_num).alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells12[10].add_paragraph(personal_tel).alignment=WD_ALIGN_PARAGRAPH.CENTER
    #
    # hdr_cells15[0].add_paragraph('\n\n本人申请理由').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells15[3].add_paragraph(reason + '\n\n学生本人签名：'+'\n\n\t\t年\t月\t日').alignment=WD_ALIGN_PARAGRAPH.CENTER
    #
    # hdr_cells21[0].add_paragraph('\n\n院系领导意见').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells21[3].add_paragraph('\n\n领导签字：'+'\n\n\t\t年\t月\t日').alignment=WD_ALIGN_PARAGRAPH.CENTER
    #
    # hdr_cells27[0].add_paragraph('\n\n学生处意见').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells27[3].add_paragraph('\n\n领导签字：'+'\n\n\t\t年\t月\t日').alignment=WD_ALIGN_PARAGRAPH.CENTER
    #
    # hdr_cells32[0].add_paragraph('\n\n主管校领导审批').alignment=WD_ALIGN_PARAGRAPH.CENTER
    # hdr_cells32[3].add_paragraph('\n\n领导签字：'+'\n\n\t\t年\t月\t日').alignment=WD_ALIGN_PARAGRAPH.CENTER
    #

    document.save(r'{}\{}来样分析表.docx'.format(url,deng_ji_hao))

    # document1 = Document()
    # document1.styles['Normal'].font.name = u'宋体'
    # document1.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # p1=document1.add_paragraph()
    # p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p1.paragraph_format.line_spacing = Pt(25)
    # r1 =p1.add_run('X X X X X X X X\rX X 清 单\n')
    # r1.font.size = Pt(16)
    # r1.bold = True

    # leave_name = document1.add_paragraph('各部门：\n\n\t')
    # leave_name.add_run(departments).font.underline = True
    # leave_name.add_run('院（系）')
    # leave_name.add_run(specialty).font.underline = True
    # leave_name.add_run('专业')
    # leave_name.add_run(class_name).font.underline = True
    # leave_name.add_run('班,学生')
    # leave_name.add_run(name).font.underline = True
    # leave_name.add_run('，（学号：')
    # leave_name.add_run(school_num).font.underline = True
    # leave_name.add_run('）。申请')
    # leave_name.add_run('XX').font.underline = True
    # leave_name.add_run(',请予以协助办理相关XX手续。')
    # leave_name.add_run('\n\n\n\n\t\t\t\t\t\t\t\t\tX X 处  （公 章）\n\n\t\t\t\t\t\t\t\t\t')
    # leave_name.add_run(nt.strftime('%Y{y}%m{m}%d{d}\n\n\n\n\n\n\n').format(y='年', m='月', d='日'))

    # table1 = document1.add_table(rows=10,cols=8,style='Table Grid')
    # table1.cell(0,0).merge(table1.cell(7,3))
    # table1.cell(0,4).merge(table1.cell(7,7))
    # table1.cell(8,0).merge(table1.cell(9,7))
    #
    # hdr1_cells0 = table1.rows[0].cells
    # hdr1_cells1 = table1.rows[8].cells
    #
    # dormitory = hdr1_cells0[0].add_paragraph()
    # dormitory.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # dormitory1 = dormitory.add_run('XXX\n\n\n（盖章）\n\n\t\t年  月  日')
    # dormitory1.bold = True
    # dormitory1.font.size = Pt(14)
    #
    # library = hdr1_cells0[4].add_paragraph()
    # library.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # library1 = library.add_run('XXX\n\n\n（盖章）\n\n\t\t年  月  日')
    # library1.bold = True
    # library1.font.size = Pt(14)
    #
    # note = hdr1_cells1[0].add_paragraph()
    # note1 = note.add_run('备注：\n')
    # note1.bold = True
    # note1.font.size = Pt(12)
    #
    #
    # document1.save(name+'XX清单.docx')
if __name__ == '__main__':
    yun_xing(r'D:\Study\pythonProject\remi_design\jisuan_shabu_canshu\210726')