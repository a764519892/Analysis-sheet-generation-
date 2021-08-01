from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
document = Document()
# document.styles['Normal'].font.name = u'宋体'
# document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
# p=document.add_paragraph()
# p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
# r =p.add_run('香港瑞升（越南）纺织品有限公司\r来 样 分 析 登 记 表')
# r.font.size = Pt(16)
# r.bold = True

for row in range(8):

    if row == 7:
        t = document.add_table(rows=1, cols=24, style='Table Grid')
        t.autofit = False  # 很重要！
        w = 0.9
        t.columns[0].width = Inches(0.5)
        t.columns[1].width = Inches(w)
    else:
        t = document.add_table(rows=1, cols=24, style='Table Grid')
        t.autofit = False  # 很重要！
        t.columns[0].width = Inches(0.5)
        t.columns[1].width = Inches(0.9)


document.save('table-step.docx')